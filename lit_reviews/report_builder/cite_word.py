

from docx import Document, styles

from pandas import read_csv

from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.dml import MSO_THEME_COLOR
from docx.shared import RGBColor


## this will hold the classes for Word Document Builders




class CiteWordDocBuilder:
    def __init__(self, output_path):

        self.word_doc = Document()

        self.styles = self.word_doc.styles
        self.output_path = output_path
        ## set Cite Styles here
        self.h1_style = self.styles.add_style("CiteH1", WD_STYLE_TYPE.PARAGRAPH)
        self.h1_style.font.name = "Arial"
        self.h1_style.font.size = Pt(30)
        self.h1_style.font.color.rgb = RGBColor.from_string("548dd4")

        h2_style = self.styles.add_style("CiteH2", WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.name = "Arial"
        h2_style.font.size = Pt(20)
        h2_style.font.color.rgb = RGBColor.from_string("548dd4")

        h3_style = self.styles.add_style("CiteH3", WD_STYLE_TYPE.PARAGRAPH)
        h3_style.font.name = "Arial"
        h3_style.font.size = Pt(14)
        h3_style.font.color.rgb = RGBColor.from_string("548dd4")

    def add_hx(self, hx_text, style):

        page_title = self.word_doc.add_paragraph()

        page_title.text = hx_text
        page_title.style = self.word_doc.styles[style]

        return True

    def add_p(self, text, style=None):

        pg = self.word_doc.add_paragraph()
        pg.text = text

        run = pg.add_run()
        run.add_break()

        ## what about hiting return?

    def add_bullets(self, bullet_list, ordered=False):

        style = "List Bullet" if ordered is False else "List Number"

        for item in bullet_list:

            self.word_doc.add_paragraph(item, style=style)

        ## TODO handle nested bullets slickly.

    def add_table_row(self, table, cell_list):

        row_cells = table.add_row().cells

        for index, val in enumerate(cell_list):
            row_cells[index].text = str(val)

    def init_table(self, header_row):

        table = self.word_doc.add_table(rows=1, cols=len(header_row))

        hdr_cells = table.rows[0].cells

        for index, col in enumerate(header_row):
            hdr_cells[index].text = col

        return table

    def save_file(self, filename):

        self.word_doc.save(self.output_path + filename)
        return self.output_path + filename


class CiteProtocolBuilder(CiteWordDocBuilder):
    def __init__(self, lit_review_id, output_path):
        self.lit_review_id = lit_review_id

        self.search_protocol = SearchProtocol.objects.get(
            literature_review__id=lit_review_id
        )

        self.device_name = "Test Device"  # get based on product ID.

        self.search_term = 10  ## need to detect this.

        # self.word_doc = Document()
        self.bp = bp  ##what is BP

        # self.styles = self.word_doc.styles

        ## set Cite Styles here
        # self.h1_style = self.styles.add_style('CiteH1', WD_STYLE_TYPE.PARAGRAPH)
        # self.h1_style.font.name = 'Arial'
        # self.h1_style.font.size = Pt(30)

        # h2_style = self.styles.add_style('CiteH2', WD_STYLE_TYPE.PARAGRAPH)
        # h2_style.font.name = 'Arial'
        # h2_style.font.size = Pt(14)

        ## How to add styles

        # self.h1_style = self.word_doc.add_style('cite h1')
        ## etc. for all styles.

        print("Running Equivalence Table Builder...")
        # self.a6_equivalence_table()
        self.a5_adverse_databases()
        self.word_doc.save(
            "./review_output/{0}/protocol.docx".format(self.lit_review_id)
        )

    def a1_background_sections(self, descriptions={}):

        ## Background
        self.add_hx("Overview", "CiteH1")

        self.add_hx("Background", "CiteH2")

        self.add_p(self.search_protocol.background)

        self.add_p(bp["overview"]["background"])

        self.add_bullets(bp["overview"]["background_bullets"])

        ## Device Description
        self.add_hx("Device Description", "CiteH2")

        self.add_p(self.search_protocol.device_description)

        ## Intended Use
        self.add_hx("Intended Use", "CiteH2")
        self.add_p(self.search_protocol.intended_use)

        ## Equivalent Devices (maybe multi paragraphs)
        self.add_hx("Equivalent Devices", "CiteH2")
        self.add_p(self.search_protocol.equivalence_discussion_clinical)
        self.add_p(self.search_protocol.equivalence_discussion_technical)
        self.add_p(self.search_protocol.equivalence_discussion_biological)

        ## Equivalence Requirements
        self.add_hx("Equivalence Requirements", "CiteH2")

        equiv_reqs_txt = bp["overview"]["equiv_reqs"].replace(
            "[DEVICENAME]", self.device_name
        )
        self.add_p(equiv_reqs_txt)
        self.add_bullets(bp["overview"]["equiv_reqs_bullets"])

    def a2_selection_criteria(
        self,
    ):

        ## intro paragraph seems too custom to do boilerplate
        self.add_hx("Literature Search Methodology and Selection Criteria", "CiteH1")

        ## Scope H2 Boilerplate + Replace term of search
        self.add_hx("Scope", "CiteH2")
        scope_txt = self.bp["criteria"]["scope"].replace(
            "[SEARCHTERM]", str(self.search_term)
        )
        self.add_p(scope_txt)

        ##Date of Search H2  (today + 1 day)
        self.add_hx("Date of Search", "CiteH2")
        today = datetime.today()
        today = today.strftime("%m/%d/%Y")
        self.add_p(today)

        ## Name of Persons H2 (boilerplate)
        self.add_hx("Name of Person Unertaking Search", "CiteH2")
        self.add_p(self.bp["criteria"]["person"])

        ## Period Covered By Search H2 boiler + replace years
        self.add_hx("Period Covered By Search")
        self.add_p(
            self.bp["criteria"]["period_covered"].replace(
                "[SEARCHTERM]", str(self.search_term)
            )
        )

        ## Scientific DBs H2
        self.add_hx("Scientific Databases", "CiteH2")
        ## Bullet list of DBs + Boilerplate
        self.add_bullets(self.bp["criteria"]["sci_dbs"])

        ## Adverse Event DBs H2
        self.add_hx("Adverse Event Databases", "CiteH2")
        ## Bullet List of DBs, boilerplate
        self.add_bullets(self.bp["criteria"]["ae_dbs"])
        ## DB SearchDetails H2 - boilerplate
        self.add_hx("Database Search Details", "CiteH2")
        self.add_p(self.bp["criteria"]["search_details"])

        ## Handling Dupes H1  - boilerplate
        self.add_hx("Handling of Duplicate Literature References", "CiteH1")
        self.add_p(self.bp["criteria"]["dupes1"])

        self.add_hx("How Duplicates Are Identified?", "CiteH3")
        self.add_p(self.bp["criteria"]["dupes2"])
        self.add_bullets(self.bp["criteria"]["dupes3_bullets"])

        ## Focused Search H1 - boilerplate
        self.add_hx("Focused Search and Abstract Review Plan", "CiteH1")

        self.add_p(self.bp["focused_search"]["fs1"])
        self.add_p(self.bp["focused_search"]["fs2"])
        self.add_p(self.bp["focused_search"]["fs3"])

        ## second paragraph to be filled out - use [REPLACE ME]

        ## Selection Criteria H1 - boilerplate
        self.add_hx("Selection Criteria", "CiteH1")
        self.add_p(self.bp["focused_search"]["select_criteria"])

        ## Inclusion Criteria H2 - boiler + replace with prod name
        self.add_hx("Inclusion Criteria", "CiteH2")
        self.add_bullets(self.bp["focused_search"]["inclusion_bullets"])
        ## bullets

        ## Exclusion Criteria H2 - boiler
        self.add_hx("Exclusion Criteria", "CiteH2")
        self.add_bullets(self.bp["focused_search"]["exclusion_bullets"])
        ## + boiler folowup paragraph
        self.add_p(self.bp["focused_search"]["exclusion"])

        # Outputs H1  - boiler
        self.add_hx("Outputs", "CiteH1")
        ## Data Selection Process
        self.add_p(self.bp["outputs"]["intro"])

        self.add_hx("Data Selection Process", "CiteH2")
        ## boiler + image

    def a3_scientific_databases(self):

        # loop each database here, grab terms
        search_props = LiteratureReviewSearchProposal.objects.filter(
            literature_review__id=self.lit_review_id,
            db__entrez_enum="pubmed",
        )
        terms = []
        for prop in search_props:

            terms.append(prop.term)

        for db in self.bp["scientific_databases"]["dbs"]:

            self.add_hx(db["name"], "CiteH2")
            self.add_p(db["link"])
            self.add_p(self.bp["scitentific_databases"]["strategy_p"])

            self.add_hx("Search Strategy", "CiteH3")
            self.add_bullets(db["strat_bullets"])

            self.add_hx("Terms", "CiteH3")
            self.add_bullets(terms)

    def a4_lit_appraisal_plan(self):

        # Title H1 + boiler paragraph
        self.add_hx("Clinical Literature Appraisal Plan", "CiteH1")
        self.add_p(self.bp["appr_plan"]["intro"])
        ## Table1 Boilerplate

        ## table 2 boilerplate

        # Clin Lit Analysis a4_lit_appraisal_plan H1

        self.add_hx("Clinical Literature Analysis Plan", "CiteH1")
        # bullets (nested if possible)
        self.add_p(self.bp["analysis_plan"]["intro"])
        self.add_bullets(self.bp["analysis_plan"]["bullets1"])

        self.add_p("Comprehensive Summary to include:")
        self.add_bullets(self.bp["analysis_plan"]["bullets2"])

    def a5_adverse_databases(self):

        self.add_hx("Appendix E Adverse Event Databases / Recalls", "CiteH1")

        # maude, recalls, mhra
        self.add_hx("US FDA MAUDE (USA)", "CiteH2")

        self.add_hx("Date Range", "CiteH2")

        self.add_hx("Search Strategy", "CiteH2")

        maude_search = AdverseEventSearch.objects.get(
            db__name="maude", literature_review__id=self.lit_review_id
        )

        self.add_p(
            self.bp["ae_maude"]["strategy_p1"].replace(
                "[REPLACECODE]", maude_search.product_code.code
            )
        )
        self.add_p(self.bp["ae_maude"]["strategy_p2"])

        self.add_hx("Adverse Event Summary", "CiteH2")

        ## summary table here.  totals.
        columns = ["Death", "Injury", "Malfunction", "Other/NA", "Excluded"]
        table = self.init_table(columns)

        death = AdverseEventReview.objects.filter(
            search=maude_search, ae__event_type="Death"
        ).count()
        injury = AdverseEventReview.objects.filter(
            search=maude_search, ae__event_type="Injury"
        ).count()
        malfunction = AdverseEventReview.objects.filter(
            search=maude_search, ae__event_type="Malfunction"
        ).count()
        other = AdverseEventReview.objects.filter(
            Q(search=maude_search)
            & Q(Q(ae__event_type="NA") | Q(ae__event_type="Other"))
        ).count()

        excluded = 0
        row = [death, injury, malfunction, other, excluded]

        self.add_table_row(table, row)

        self.add_p("")

        ## summary table breakdown by year.
        columns = ["Code", "Year", "Number of Events", "Number of Related Events"]
        table = self.init_table(columns)
        for i in range(0, self.search_term):
            this_year = datetime.today()

            this_year = (this_year - timedelta(days=365 * i)).year
            ae_reviews_total = AdverseEventReview.objects.filter(
                search=maude_search, ae__event_date__year=this_year
            )

            ae_reviews_relevant = AdverseEventReview.objects.filter(
                search=maude_search, ae__event_date__year=this_year, state="IN"
            )
            row = [
                maude_search.product_code.code,
                this_year,
                len(ae_reviews_total),
                len(ae_reviews_relevant),
            ]
            self.add_table_row(table, row)
